from typing import Any, Dict, Optional
import logging
import io

import fitz  # PyMuPDF
import docx  # python-docx
from fastapi import APIRouter, Body, File, Form, HTTPException, UploadFile
from google.api_core import exceptions as google_api_exceptions

from services.graphrag import NoGeminiModelAvailable, run_graphrag_pipeline

logger = logging.getLogger(__name__)
router = APIRouter()

ALLOWED_EXTENSIONS = (".pdf", ".docx", ".doc")


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text content from a PDF file using PyMuPDF."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []
        for page in doc:
            pages_text.append(page.get_text())
        doc.close()
        return "\n".join(pages_text).strip()
    except Exception as exc:
        raise ValueError(f"PDF dosyası okunamadı: {exc}") from exc


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text content from a Word (.docx) file using python-docx."""
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except Exception as exc:
        raise ValueError(f"Word dosyası okunamadı: {exc}") from exc


def _extract_text_from_file(filename: str, file_bytes: bytes) -> str:
    """Route to the correct extractor based on file extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_text_from_pdf(file_bytes)
    elif lower.endswith(".docx") or lower.endswith(".doc"):
        return _extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Desteklenmeyen dosya formatı: {filename}")


from fastapi import APIRouter, Body, File, Form, HTTPException, UploadFile, Request

@router.post("/report")
async def generate_report(request: Request):
    content_type = request.headers.get("content-type", "")
    
    system_name = None
    text = None
    files = []
    
    if "multipart/form-data" in content_type:
        form = await request.form()
        system_name = form.get("system_name")
        text = form.get("text")
        
        files = form.getlist("file")
        if not files and form.get("files"):
            files = form.getlist("files")
    else:
        # JSON body
        try:
            body = await request.json()
            system_name = body.get("system_name")
            text = body.get("text")
        except Exception:
            pass
    file_texts = []
    for file in files:
        if hasattr(file, "read") and file.filename:
            if not any(file.filename.lower().endswith(ext) for ext in ALLOWED_EXTENSIONS):
                raise HTTPException(
                    status_code=422,
                    detail=f"Sadece PDF ve Word (.docx) dosyaları desteklenmektedir. Bulunan: {file.filename}",
                )
            file_bytes = await file.read()
            if not file_bytes:
                continue
            extracted = _extract_text_from_file(file.filename, file_bytes)
            if extracted.strip():
                file_texts.append(extracted)

    # Combine manual text + file texts
    combined_text = "\n\n".join(
        part for part in [text or ""] + file_texts if part.strip()
    )

    effective_system_name = system_name or None

    if not effective_system_name and not combined_text.strip():
        raise HTTPException(
            status_code=422,
            detail="En az bir system_name, metin veya dosya sağlanmalıdır.",
        )

    logger.info(
        "POST /report called for system: %s | text_len=%d | file_count=%d",
        effective_system_name or "Raw Text",
        len(combined_text),
        len(files),
    )

    # If the user is just requesting an existing system's report from history (no new text/file),
    # we should return the exact saved report instead of regenerating it with the LLM.
    if effective_system_name and not combined_text.strip():
        from db.systems import get_system_saved_report
        saved_report = get_system_saved_report(effective_system_name)
        if saved_report:
            logger.info(f"Returning previously saved report for system: {effective_system_name}")
            return {
                "system": effective_system_name,
                "report": saved_report
            }

    try:
        result = run_graphrag_pipeline(effective_system_name, combined_text)
    except ValueError as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc
    except NoGeminiModelAvailable as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc
    except google_api_exceptions.InvalidArgument as exc:
        err = str(exc).lower()
        if "api key" in err or "api_key" in err:
            raise HTTPException(
                status_code=502,
                detail=(
                    "Gemini API anahtarı geçersiz veya eksik. "
                    ".env içinde GEMINI_API_KEY değerini Google AI Studio'dan aldığın geçerli bir anahtarla güncelle; "
                    "uvicorn'u durdurup yeniden başlat."
                ),
            ) from exc
        raise HTTPException(status_code=502, detail=str(exc)) from exc
    except google_api_exceptions.GoogleAPIError as exc:
        raise HTTPException(
            status_code=502,
            detail=f"Gemini API hatası: {exc}",
        ) from exc

    # If the user provided a system name (or one was generated), save the new report scores to the Neo4j DB
    sys_name = result.get("system", effective_system_name)
    if sys_name and sys_name != "Unknown System":
        from db.systems import save_system_report_to_db
        try:
            save_system_report_to_db(sys_name, result.get("report", {}))
            logger.info(f"Saved generated risk scores for system '{sys_name}' to Neo4j.")
        except Exception as e:
            logger.error(f"Failed to save system report to DB: {e}")

    return result
