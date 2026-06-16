from typing import Optional
import io

# pyrefly: ignore [missing-import]
import fitz  # PyMuPDF
import docx  # python-docx
from fastapi import APIRouter, Body, File, Form, HTTPException, Request, UploadFile
import logging
from services.analysis_service import analyze_text, generate_graph_trace
from models.schemas import AnalyzeTextResponse

logger = logging.getLogger(__name__)
router = APIRouter()

ALLOWED_EXTENSIONS = (".pdf", ".docx", ".doc")


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text content from a PDF file using PyMuPDF."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []
        for page in doc:
            pages_text.append(page.get_text())
        doc.close()
        return "\n".join(pages_text).strip()
    except Exception as exc:
        raise ValueError(f"PDF dosyası okunamadı: {exc}") from exc


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text content from a Word (.docx) file using python-docx."""
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except Exception as exc:
        raise ValueError(f"Word dosyası okunamadı: {exc}") from exc


def _extract_text_from_file(filename: str, file_bytes: bytes) -> str:
    """Route to the correct extractor based on file extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_text_from_pdf(file_bytes)
    elif lower.endswith(".docx") or lower.endswith(".doc"):
        return _extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Desteklenmeyen dosya formatı: {filename}")


@router.post("/analyze-text", response_model=AnalyzeTextResponse)
async def analyze_text_endpoint(request: Request):
    """
    Analyze free text (and/or PDF/Word) against keyword mappings stored in Neo4j.

    Accepts either:
      - JSON body: {"text": "..."}
      - multipart/form-data with text (optional) + file (optional PDF/DOCX)

    Returns:
      - matched_keywords, inferred_system_types, inferred_risks, inferred_regulations
    """
    content_type = request.headers.get("content-type", "")

    if "multipart/form-data" in content_type:
        form = await request.form()
        text = form.get("text", "") or ""
        files = form.getlist("file")
        if not files and form.get("files"):
            files = form.getlist("files")

        file_texts = []
        for file in files:
            if hasattr(file, "read") and file.filename:
                if not any(file.filename.lower().endswith(ext) for ext in ALLOWED_EXTENSIONS):
                    raise HTTPException(
                        status_code=422,
                        detail=f"Sadece PDF ve Word (.docx) dosyaları desteklenmektedir. Bulunan: {file.filename}",
                    )
                file_bytes = await file.read()
                if file_bytes:
                    file_texts.append(_extract_text_from_file(file.filename, file_bytes))

        combined_text = "\n\n".join(
            part for part in [text] + file_texts if part.strip()
        )
    else:
        # JSON body
        body = await request.json()
        combined_text = body.get("text", "")

    if not combined_text.strip():
        raise HTTPException(
            status_code=422, detail="Metin veya dosya sağlanmalıdır."
        )

    logger.info("POST /analyze-text called (text_len=%d)", len(combined_text))
    result = analyze_text(combined_text)
    return result


@router.post("/graph-trace")
async def graph_trace_endpoint(request: Request):
    """
    Returns an explainable reasoning chain for ontology inference.
    Deterministic, no LLM. Accepts JSON or multipart/form-data.
    """
    content_type = request.headers.get("content-type", "")

    if "multipart/form-data" in content_type:
        form = await request.form()
        text = form.get("text", "") or ""
        files = form.getlist("file")
        if not files and form.get("files"):
            files = form.getlist("files")

        file_texts = []
        for file in files:
            if hasattr(file, "read") and file.filename:
                if not any(file.filename.lower().endswith(ext) for ext in ALLOWED_EXTENSIONS):
                    raise HTTPException(
                        status_code=422,
                        detail=f"Sadece PDF ve Word (.docx) dosyaları desteklenmektedir. Bulunan: {file.filename}",
                    )
                file_bytes = await file.read()
                if file_bytes:
                    file_texts.append(_extract_text_from_file(file.filename, file_bytes))

        combined_text = "\n\n".join(
            part for part in [text] + file_texts if part.strip()
        )
    else:
        # JSON body
        body = await request.json()
        combined_text = body.get("text", "")

    logger.info("POST /graph-trace called (text_len=%d)", len(combined_text))
    return generate_graph_trace(combined_text)
